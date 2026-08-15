import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import numpy as np
import time
import os
import glob
import base64
import plotly.graph_objects as go
import plotly.express as px
from sklearn.ensemble import RandomForestRegressor
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
import lightgbm as lgb
import io

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

def find_asset(basename):
    """Cari file di folder script berdasarkan nama saja, tidak peduli ekstensinya
    (mis. banner_kopi.jpg / .jpeg / .png / .webp semua akan ketemu)."""
    matches = sorted(glob.glob(os.path.join(SCRIPT_DIR, basename + ".*")))
    return matches[0] if matches else None

FAVICON = find_asset("favicon")
BANNER_LOGIN = find_asset("preview_banner_dashboard")
BANNER_MAIN = find_asset("preview_banner_dashboard")

st.set_page_config(
    page_title="Dashboard Prediksi Pendapatan - Sumatra Roastery Medan",
    page_icon=FAVICON if FAVICON else "☕",
    layout="wide",
)

PRIMARY = "#0F6B5C"
ACCENT = "#B5502D"
GRID = "#e5e0d8"

CREAM = "#FAF6F0"
CREAM_DARK = "#F0E8DC"
ESPRESSO = "#3B2A20"
ESPRESSO_SOFT = "#5C4633"
BORDER = "#D9C9B4"
GOLD = "#E0A526"

st.markdown(f"""
<style>
/* ---- Base app background & typography ---- */
.stApp {{
    background: linear-gradient(135deg, #EAF2EF 0%, {CREAM} 45%, #FBEEE4 100%) !important;
    background-attachment: fixed;
}}
html, body, [class*="css"] {{
    color: {ESPRESSO} !important;
    font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
}}

/* ---- Kill the stray white top toolbar ---- */
[data-testid="stHeader"] {{
    background-color: transparent !important;
}}
[data-testid="stToolbar"] {{
    background-color: transparent;
}}

/* ---- Naikkan konten utama supaya sejajar dengan bagian atas sidebar ---- */
[data-testid="stMainBlockContainer"], .block-container {{
    padding-top: 1.5rem !important;
}}

/* ---- Sidebar ---- */
[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, #EAF2EF 0%, {CREAM_DARK} 55%, #FBEEE4 100%) !important;
    border-right: 1px solid {BORDER};
}}
[data-testid="stSidebar"] * {{
    color: {ESPRESSO} !important;
}}

/* ---- Headings, captions, labels, markdown text ---- */
h1, h2, h3, h4, h5, h6, p, span, label, .stMarkdown, .stCaption {{
    color: {ESPRESSO} !important;
}}
[data-testid="stCaptionContainer"] {{
    color: {ESPRESSO_SOFT} !important;
}}

/* ---- Text inputs & password fields ---- */
.stTextInput input, .stNumberInput input {{
    background-color: #FFFFFF !important;
    color: {ESPRESSO} !important;
    border: 1px solid {BORDER} !important;
    border-radius: 6px;
}}
.stTextInput input::placeholder {{
    color: #A8998A !important;
}}
.stTextInput label, .stNumberInput label, .stSlider label, .stSelectbox label, .stFileUploader label {{
    color: {ESPRESSO} !important;
    font-weight: 600;
}}

/* ---- Selectbox ---- */
.stSelectbox div[data-baseweb="select"] > div {{
    background-color: #FFFFFF !important;
    color: {ESPRESSO} !important;
    border: 1px solid {BORDER} !important;
}}

/* ---- Buttons (target inner text explicitly so it doesn't get overridden) ---- */
.stButton button, .stFormSubmitButton button {{
    background-color: {ESPRESSO} !important;
    border: none !important;
    border-radius: 6px;
    font-weight: 600;
}}
.stButton button p, .stFormSubmitButton button p,
.stButton button div, .stFormSubmitButton button div,
.stButton button span, .stFormSubmitButton button span {{
    color: {CREAM} !important;
}}
.stButton button:hover, .stFormSubmitButton button:hover {{
    background-color: {ESPRESSO_SOFT} !important;
}}

/* ---- Neutralize Chrome/browser autofill styling on inputs ---- */
input:-webkit-autofill,
input:-webkit-autofill:hover,
input:-webkit-autofill:focus {{
    -webkit-box-shadow: 0 0 0px 1000px #FFFFFF inset !important;
    -webkit-text-fill-color: {ESPRESSO} !important;
    caret-color: {ESPRESSO} !important;
}}

/* ---- Banner images: cap height so they never dominate the page ---- */
[data-testid="stImage"] img {{
    max-height: 300px;
    width: 100%;
    object-fit: cover;
    border-radius: 18px;
}}

/* ---- Slider track/thumb ---- */
[data-testid="stSlider"] [role="slider"] {{
    background-color: {PRIMARY} !important;
}}

/* ---- File uploader box ---- */
[data-testid="stFileUploaderDropzone"] {{
    background-color: #FFFFFF !important;
    border: 1.5px dashed {BORDER} !important;
}}
[data-testid="stFileUploaderDropzone"] * {{
    color: {ESPRESSO_SOFT} !important;
}}

/* ---- Tabs ---- */
.stTabs [data-baseweb="tab-list"] {{
    gap: 6px;
    border-bottom: 3px solid {BORDER};
}}
.stTabs [data-baseweb="tab"],
.stTabs [role="tab"] {{
    color: {ESPRESSO_SOFT} !important;
    background: transparent !important;
    border-radius: 999px !important;
    padding: 6px 16px !important;
    transition: background 0.15s ease, color 0.15s ease, box-shadow 0.15s ease;
}}
.stTabs [data-baseweb="tab"] *,
.stTabs [role="tab"] * {{
    font-weight: 700 !important;
    font-size: 1rem !important;
}}
.stTabs [data-baseweb="tab"]:hover,
.stTabs [role="tab"]:hover {{
    color: {ESPRESSO} !important;
}}
.stTabs [aria-selected="true"] {{
    color: #FFFFFF !important;
    background: linear-gradient(135deg, #4FB6E8 0%, #2E93C9 100%) !important;
    box-shadow: 0 3px 8px rgba(47, 147, 201, 0.35);
    border-bottom: none !important;
}}

/* ---- Metrics ---- */
[data-testid="stMetric"] {{
    background-color: #FFFFFF;
    border: 1px solid {BORDER};
    border-radius: 8px;
    padding: 14px 16px;
}}
[data-testid="stMetricLabel"] {{
    color: {ESPRESSO_SOFT} !important;
}}
[data-testid="stMetricValue"] {{
    color: {ESPRESSO} !important;
}}

/* ---- Dataframes ---- */
[data-testid="stDataFrame"] {{
    border: 1px solid {BORDER};
    border-radius: 6px;
}}

/* ---- Expander ---- */
[data-testid="stExpander"] {{
    background-color: #FFFFFF;
    border: 1px solid {BORDER};
    border-radius: 8px;
}}
</style>
""", unsafe_allow_html=True)

BULAN_ORDER = ['Januari','Februari','Maret','April','Mei','Juni','Juli','Agustus','September','Oktober','November','Desember']
BULAN_MAP = {b: i + 1 for i, b in enumerate(BULAN_ORDER)}


@st.cache_data
def load_raw(file):
    xls = pd.ExcelFile(file)
    daily = pd.read_excel(xls, sheet_name='Dataset Harian 2023-2025', header=3)
    per_jenis = pd.read_excel(xls, sheet_name='Pendapatan Per Jenis Kopi', header=2)
    rekap = pd.read_excel(xls, sheet_name='Rekap Bulanan', header=2)
    return daily, per_jenis, rekap


@st.cache_data
def build_dataset(daily, per_jenis, rekap):
    agg_harga = (
        daily.groupby(['Tahun', 'Bulan', 'Jenis Kopi'])
        .apply(lambda g: (g['Harga (Rp)'] * g['Jumlah Terjual']).sum() / g['Jumlah Terjual'].sum()
               if g['Jumlah Terjual'].sum() > 0 else g['Harga (Rp)'].mean())
        .reset_index(name='harga_rata2')
    )

    df = per_jenis.merge(agg_harga, on=['Tahun', 'Bulan', 'Jenis Kopi'], how='left')
    df['bulan_num'] = df['Bulan'].map(BULAN_MAP)
    df['periode'] = df['Tahun'] * 100 + df['bulan_num']
    df = df.sort_values(['Jenis Kopi', 'periode']).reset_index(drop=True)
    df['lag_1'] = df.groupby('Jenis Kopi')['Total Pendapatan (Rp)'].shift(1)

    rekap = rekap.copy()
    rekap['bulan_num'] = rekap['Bulan'].map(BULAN_MAP)
    rekap['periode'] = rekap['Tahun'] * 100 + rekap['bulan_num']
    avg_overall = rekap['Total Pendapatan (Rp)'].mean()
    rekap['kategori_tren'] = np.where(rekap['Total Pendapatan (Rp)'] >= avg_overall, 'Tinggi', 'Rendah')

    df = df.merge(rekap[['periode', 'kategori_tren']], on='periode', how='left')
    return df, rekap, avg_overall


@st.cache_data
def train_models(df, split_ratio):
    le_jenis = LabelEncoder()
    le_tren = LabelEncoder()
    df_model = df.dropna(subset=['lag_1']).copy()
    df_model['jenis_kopi_enc'] = le_jenis.fit_transform(df_model['Jenis Kopi'])
    df_model['kategori_tren_enc'] = le_tren.fit_transform(df_model['kategori_tren'])
    df_model = df_model.sort_values('periode').reset_index(drop=True)

    features = ['Tahun', 'bulan_num', 'jenis_kopi_enc', 'harga_rata2', 'lag_1', 'kategori_tren_enc']
    target = 'Total Pendapatan (Rp)'

    periodes = sorted(df_model['periode'].unique())
    split_periode = periodes[int(len(periodes) * split_ratio)]
    train = df_model[df_model['periode'] < split_periode]
    test = df_model[df_model['periode'] >= split_periode]

    X_train, y_train = train[features], train[target]
    X_test, y_test = test[features], test[target]

    results = {}
    preds = {}

    t0 = time.time()
    rf = RandomForestRegressor(n_estimators=200, random_state=42)
    rf.fit(X_train, y_train)
    rf_time = time.time() - t0
    pred_rf = rf.predict(X_test)
    results['Random Forest'] = dict(
        MAE=mean_absolute_error(y_test, pred_rf),
        RMSE=np.sqrt(mean_squared_error(y_test, pred_rf)),
        R2=r2_score(y_test, pred_rf),
        Training_Time=rf_time,
    )
    preds['Random Forest'] = pred_rf

    t0 = time.time()
    lgbm = lgb.LGBMRegressor(n_estimators=200, random_state=42, verbose=-1)
    lgbm.fit(X_train, y_train)
    lgb_time = time.time() - t0
    pred_lgb = lgbm.predict(X_test)
    results['LightGBM'] = dict(
        MAE=mean_absolute_error(y_test, pred_lgb),
        RMSE=np.sqrt(mean_squared_error(y_test, pred_lgb)),
        R2=r2_score(y_test, pred_lgb),
        Training_Time=lgb_time,
    )
    preds['LightGBM'] = pred_lgb

    fi = pd.DataFrame({
        'Fitur': features,
        'Random Forest': rf.feature_importances_ / rf.feature_importances_.sum(),
        'LightGBM': lgbm.feature_importances_ / lgbm.feature_importances_.sum(),
    })

    test_out = test[['Tahun', 'Bulan', 'Jenis Kopi', 'Total Pendapatan (Rp)']].copy()
    test_out['Prediksi Random Forest'] = pred_rf
    test_out['Prediksi LightGBM'] = pred_lgb

    return results, fi, test_out, split_periode


def rupiah(x):
    return f"Rp {x:,.0f}".replace(",", ".")


def to_excel_bytes(df, sheet_name="Data"):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return buf.getvalue()


def generate_laporan_docx(periode_text, jenis_text, total_pendapatan, total_qty, rata_harga, ringkasan_jenis):
    doc = Document()
    doc.add_heading('Laporan Penjualan — Sumatra Roastery Medan', level=1)
    doc.add_paragraph(f"Periode: {periode_text}")
    doc.add_paragraph(f"Jenis Kopi: {jenis_text}")
    doc.add_paragraph(f"Tanggal Cetak: {pd.Timestamp.now().strftime('%d %B %Y')}")

    doc.add_heading('Ringkasan', level=2)
    doc.add_paragraph(f"Total Pendapatan: {rupiah(total_pendapatan)}")
    doc.add_paragraph(f"Total Unit Terjual: {int(total_qty):,}".replace(",", "."))
    doc.add_paragraph(f"Rata-rata Harga: {rupiah(rata_harga)}")

    doc.add_heading('Rincian per Jenis Kopi', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Jenis Kopi', 'Total Pendapatan (Rp)', 'Jumlah Terjual'
    for _, row in ringkasan_jenis.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Jenis Kopi'])
        cells[1].text = rupiah(row['Total Pendapatan (Rp)'])
        cells[2].text = str(int(row['Jumlah Terjual']))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_laporan_pdf(periode_text, jenis_text, total_pendapatan, total_qty, rata_harga, ringkasan_jenis):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("Laporan Penjualan — Sumatra Roastery Medan", styles['Title']),
        Spacer(1, 12),
        Paragraph(f"Periode: {periode_text}", styles['Normal']),
        Paragraph(f"Jenis Kopi: {jenis_text}", styles['Normal']),
        Paragraph(f"Tanggal Cetak: {pd.Timestamp.now().strftime('%d %B %Y')}", styles['Normal']),
        Spacer(1, 12),
        Paragraph("Ringkasan", styles['Heading2']),
        Paragraph(f"Total Pendapatan: {rupiah(total_pendapatan)}", styles['Normal']),
        Paragraph(f"Total Unit Terjual: {int(total_qty):,}".replace(",", "."), styles['Normal']),
        Paragraph(f"Rata-rata Harga: {rupiah(rata_harga)}", styles['Normal']),
        Spacer(1, 12),
        Paragraph("Rincian per Jenis Kopi", styles['Heading2']),
        Spacer(1, 6),
    ]

    data = [['Jenis Kopi', 'Total Pendapatan (Rp)', 'Jumlah Terjual']]
    for _, row in ringkasan_jenis.iterrows():
        data.append([row['Jenis Kopi'], rupiah(row['Total Pendapatan (Rp)']), str(int(row['Jumlah Terjual']))])
    tbl = Table(data, hAlign='LEFT')
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(PRIMARY)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
    ]))
    elements.append(tbl)
    doc.build(elements)
    return buf.getvalue()


@st.cache_data
def forecast_next_month(df):
    """Latih model pakai SELURUH data (bukan hanya data training) lalu prediksi bulan setelah periode terakhir."""
    le_jenis = LabelEncoder()
    le_tren = LabelEncoder()
    df_model = df.dropna(subset=['lag_1']).copy()
    df_model['jenis_kopi_enc'] = le_jenis.fit_transform(df_model['Jenis Kopi'])
    df_model['kategori_tren_enc'] = le_tren.fit_transform(df_model['kategori_tren'])
    df_model = df_model.sort_values('periode').reset_index(drop=True)

    features = ['Tahun', 'bulan_num', 'jenis_kopi_enc', 'harga_rata2', 'lag_1', 'kategori_tren_enc']
    target = 'Total Pendapatan (Rp)'

    rf_full = RandomForestRegressor(n_estimators=200, random_state=42)
    rf_full.fit(df_model[features], df_model[target])
    lgb_full = lgb.LGBMRegressor(n_estimators=200, random_state=42, verbose=-1)
    lgb_full.fit(df_model[features], df_model[target])

    last_periode = df_model['periode'].max()
    last_tahun, last_bulan_num = last_periode // 100, last_periode % 100
    next_bulan_num = 1 if last_bulan_num == 12 else last_bulan_num + 1
    next_tahun = last_tahun + 1 if last_bulan_num == 12 else last_tahun
    next_bulan_nama = BULAN_ORDER[next_bulan_num - 1]

    last_kategori_tren = df_model.loc[df_model['periode'] == last_periode, 'kategori_tren_enc'].iloc[0]

    rows = []
    for jenis in df_model['Jenis Kopi'].unique():
        sub = df_model[df_model['Jenis Kopi'] == jenis].sort_values('periode')
        lag_1_next = sub[target].iloc[-1]
        harga_next = sub['harga_rata2'].tail(3).mean()
        jenis_enc = sub['jenis_kopi_enc'].iloc[-1]
        x_next = pd.DataFrame([{
            'Tahun': next_tahun, 'bulan_num': next_bulan_num, 'jenis_kopi_enc': jenis_enc,
            'harga_rata2': harga_next, 'lag_1': lag_1_next, 'kategori_tren_enc': last_kategori_tren,
        }])[features]
        pred_rf = rf_full.predict(x_next)[0]
        pred_lgb = lgb_full.predict(x_next)[0]
        rows.append({'Jenis Kopi': jenis, 'Pendapatan Bulan Terakhir (Rp)': lag_1_next,
                      'Prediksi Random Forest (Rp)': pred_rf, 'Prediksi LightGBM (Rp)': pred_lgb})

    forecast_df = pd.DataFrame(rows).sort_values('Prediksi Random Forest (Rp)', ascending=False).reset_index(drop=True)
    return forecast_df, next_bulan_nama, next_tahun


# ---------- LOGIN DUA PERAN ----------
from login import require_login

require_login({
    "PRIMARY": PRIMARY,
    "ACCENT": ACCENT,
    "CREAM": CREAM,
    "ESPRESSO": ESPRESSO,
    "ESPRESSO_SOFT": ESPRESSO_SOFT,
    "BORDER": BORDER,
})
@st.cache_data
def get_base64_image(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


if BANNER_MAIN:
    photo_b64 = get_base64_image(BANNER_MAIN)
    photo_ext = os.path.splitext(BANNER_MAIN)[1].replace(".", "")
    if photo_ext == "jpg":
        photo_ext = "jpeg"

    BANNER_HEIGHT = 360

    banner_html = f"""
    <html>
    <head>
    <style>
        html, body {{
            margin:0;
            padding:0;
            width:100%;
            background:transparent;
            overflow:hidden;
        }}
        .hero-banner{{
            position:relative;
            width:100%;
            height:{BANNER_HEIGHT}px;
            border-radius:18px;
            overflow:hidden;
            background:#0d3b2e;
            box-sizing:border-box;
        }}
        .hero-image{{
            position:absolute;
            left:0;
            top:0;
            width:100%;
            height:100%;
            object-fit:cover;
            object-position:center;
            display:block;
        }}
    </style>
    </head>
    <body>
        <div class="hero-banner">
            <img class="hero-image" src="data:image/{photo_ext};base64,{photo_b64}">
        </div>
    </body>
    </html>
    """

    components.html(banner_html, height=BANNER_HEIGHT, scrolling=False)



else:
    st.title("☕ Dashboard Analisis Tren & Prediksi Pendapatan")
    st.caption("Sumatra Roastery Medan — Random Forest vs LightGBM")
    st.warning(
        "File banner `preview_banner_dashboard.*` tidak ditemukan di folder aplikasi "
        f"({SCRIPT_DIR}). Pastikan file gambar ini sudah di-upload/push ke repository "
        "yang sama dengan app.py, sejajar (bukan di dalam subfolder lain)."
    )

with st.sidebar:
    if st.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.role = None
        st.rerun()
    st.divider()

IS_PENELITI = st.session_state.role == "Peneliti"

with st.sidebar:
    st.header("Data & Pengaturan")
    uploaded = None
    split_ratio = 0.8
    if IS_PENELITI:
        uploaded = st.file_uploader("Unggah dataset (.xlsx)", type=["xlsx"])
        split_ratio = st.slider("Proporsi data training", 0.6, 0.9, 0.8, 0.05)
    else:
        st.caption("Data & pengaturan hanya dapat diubah oleh Peneliti. Anda melihat hasil analisis terkini.")
    st.caption("Sesuai BAB III: time-based split 80:20")

def find_local_dataset():
    candidates = glob.glob(os.path.join(SCRIPT_DIR, "*.xlsx"))
    return candidates[0] if candidates else None

data_path = uploaded if uploaded is not None else find_local_dataset()

if data_path is None:
    st.warning("Tidak ada file .xlsx ditemukan di folder yang sama dengan app.py. Unggah dataset lewat sidebar di kiri.")
    st.stop()

try:
    daily, per_jenis, rekap_raw = load_raw(data_path)
except Exception as e:
    st.error(f"Gagal membaca dataset: {e}")
    st.stop()

df, rekap, avg_overall = build_dataset(daily, per_jenis, rekap_raw)
results, fi, test_out, split_periode = train_models(df, split_ratio)
forecast_df, next_bulan_nama, next_tahun = forecast_next_month(df)

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📋 Data Aktual", "📈 Analisis Tren", "🔮 Prediksi & Evaluasi", "⭐ Feature Importance", "📅 Prediksi Bulan Depan", "📄 Laporan"])

components.html("""
<script>
function attachTabAutoScroll() {
    try {
        const doc = window.parent.document;
        const containers = doc.querySelectorAll('[data-testid="stTabs"]');
        containers.forEach(function(container) {
            if (container.dataset.autoscrollBound) return;
            container.dataset.autoscrollBound = "1";
            container.addEventListener('click', function(e) {
                setTimeout(function() {
                    const rect = container.getBoundingClientRect();
                    const y = rect.top + doc.defaultView.scrollY - 70;
                    doc.defaultView.scrollTo({top: y, behavior: 'smooth'});
                }, 80);
            });
        });
    } catch (err) {
        // akses cross-origin diblokir browser, abaikan
    }
}
attachTabAutoScroll();
setInterval(attachTabAutoScroll, 800);
</script>
""", height=1)

with tab1:
    st.subheader("Rekap Pendapatan Bulanan")
    rekap_show = rekap[['Tahun', 'Bulan', 'Total Pendapatan (Rp)', 'kategori_tren']].copy()
    rekap_show['Total Pendapatan (Rp)'] = rekap_show['Total Pendapatan (Rp)'].apply(rupiah)
    rekap_show = rekap_show.rename(columns={'kategori_tren': 'Kategori Tren'})
    st.dataframe(rekap_show, use_container_width=True, hide_index=True)

    st.subheader("Pendapatan per Jenis Kopi per Bulan")
    per_jenis_show = per_jenis.copy()
    per_jenis_show['Total Pendapatan (Rp)'] = per_jenis_show['Total Pendapatan (Rp)'].apply(rupiah)
    per_jenis_show['% dari Total Bulan'] = (per_jenis_show['% dari Total Bulan'] * 100).round(2).astype(str) + '%'
    st.dataframe(per_jenis_show, use_container_width=True, hide_index=True)

    st.subheader("Data Transaksi Harian")
    daily_show = daily.head(50).copy()
    daily_show['Harga (Rp)'] = daily_show['Harga (Rp)'].apply(rupiah)
    daily_show['Pendapatan (Rp)'] = daily_show['Pendapatan (Rp)'].apply(rupiah)
    st.dataframe(daily_show, use_container_width=True, hide_index=True)
    st.download_button("⬇️ Download Excel — Data Transaksi Harian",
                        data=to_excel_bytes(daily_show, "Transaksi Harian"),
                        file_name="data_transaksi_harian.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

with tab2:
    st.subheader("Tren Pendapatan Bulanan (2023–2025)")
    labels = [f"{b[:3]} {t}" for b, t in zip(rekap['Bulan'], rekap['Tahun'])]
    nilai_tinggi = np.where(rekap['kategori_tren'] == 'Tinggi', rekap['Total Pendapatan (Rp)'], np.nan)
    nilai_rendah = np.where(rekap['kategori_tren'] == 'Rendah', rekap['Total Pendapatan (Rp)'], np.nan)

    fig = go.Figure()
    fig.add_trace(go.Bar(x=labels, y=nilai_tinggi, marker_color=PRIMARY,
                          name="Tinggi (≥ rata-rata)"))
    fig.add_trace(go.Bar(x=labels, y=nilai_rendah, marker_color=ACCENT,
                          name="Rendah (< rata-rata)"))
    fig.add_hline(y=avg_overall, line_dash="dash", line_color="#888",
                  annotation_text=f"Rata-rata: {rupiah(avg_overall)}")
    fig.update_layout(height=460, plot_bgcolor="white", yaxis_title="Pendapatan (Rp)",
                       xaxis_tickangle=-60, showlegend=True,
                       legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
    st.plotly_chart(fig, use_container_width=True)

    c1, c2, c3 = st.columns(3)
    c1.metric("Rata-rata Pendapatan Bulanan", rupiah(avg_overall))
    c2.metric("Bulan Kategori Tinggi", int((rekap['kategori_tren'] == 'Tinggi').sum()))
    c3.metric("Bulan Kategori Rendah", int((rekap['kategori_tren'] == 'Rendah').sum()))

    st.subheader("Pendapatan per Jenis Kopi (Total 2023–2025)")
    jenis_total = per_jenis.groupby('Jenis Kopi')['Total Pendapatan (Rp)'].sum().sort_values(ascending=False).reset_index()
    fig2 = px.bar(jenis_total, x='Jenis Kopi', y='Total Pendapatan (Rp)', color_discrete_sequence=[PRIMARY])
    fig2.update_layout(height=380, plot_bgcolor="white")
    st.plotly_chart(fig2, use_container_width=True)

with tab3:
    st.subheader("🔮 Prediksi & Evaluasi Model")
    st.caption("Hasil prediksi tiap model ditampilkan terpisah, lalu dibandingkan performanya di sub-tab terakhir.")

    subtab_rf, subtab_lgb, subtab_perf = st.tabs(["🌲 Random Forest", "💡 LightGBM", "📊 Perbandingan Performa"])

    # ---------- Sub-tab: Random Forest ----------
    with subtab_rf:
        st.markdown("#### Hasil Prediksi Random Forest — Data Uji")
        jenis_rf = st.selectbox("Pilih jenis kopi", sorted(test_out['Jenis Kopi'].unique()), key="jenis_rf")
        sub_rf = test_out[test_out['Jenis Kopi'] == jenis_rf].copy()
        sub_rf['label'] = sub_rf['Bulan'].str[:3] + " " + sub_rf['Tahun'].astype(str)

        fig_rf = go.Figure()
        fig_rf.add_trace(go.Scatter(x=sub_rf['label'], y=sub_rf['Total Pendapatan (Rp)'], name="Aktual",
                                     mode='lines+markers', line=dict(color="#333", width=3)))
        fig_rf.add_trace(go.Scatter(x=sub_rf['label'], y=sub_rf['Prediksi Random Forest'], name="Prediksi Random Forest",
                                     mode='lines+markers', line=dict(color=PRIMARY, dash='dash')))
        fig_rf.update_layout(height=420, plot_bgcolor="white", yaxis_title="Pendapatan (Rp)")
        st.plotly_chart(fig_rf, use_container_width=True)

        with st.expander("Lihat tabel hasil prediksi Random Forest"):
            rf_show = test_out[['Tahun', 'Bulan', 'Jenis Kopi', 'Total Pendapatan (Rp)', 'Prediksi Random Forest']].copy()
            rf_show['Total Pendapatan (Rp)'] = rf_show['Total Pendapatan (Rp)'].apply(rupiah)
            rf_show['Prediksi Random Forest'] = rf_show['Prediksi Random Forest'].apply(rupiah)
            st.dataframe(rf_show, use_container_width=True, hide_index=True)

    # ---------- Sub-tab: LightGBM ----------
    with subtab_lgb:
        st.markdown("#### Hasil Prediksi LightGBM — Data Uji")
        jenis_lgb = st.selectbox("Pilih jenis kopi", sorted(test_out['Jenis Kopi'].unique()), key="jenis_lgb")
        sub_lgb = test_out[test_out['Jenis Kopi'] == jenis_lgb].copy()
        sub_lgb['label'] = sub_lgb['Bulan'].str[:3] + " " + sub_lgb['Tahun'].astype(str)

        fig_lgb = go.Figure()
        fig_lgb.add_trace(go.Scatter(x=sub_lgb['label'], y=sub_lgb['Total Pendapatan (Rp)'], name="Aktual",
                                      mode='lines+markers', line=dict(color="#333", width=3)))
        fig_lgb.add_trace(go.Scatter(x=sub_lgb['label'], y=sub_lgb['Prediksi LightGBM'], name="Prediksi LightGBM",
                                      mode='lines+markers', line=dict(color=ACCENT, dash='dot')))
        fig_lgb.update_layout(height=420, plot_bgcolor="white", yaxis_title="Pendapatan (Rp)")
        st.plotly_chart(fig_lgb, use_container_width=True)

        with st.expander("Lihat tabel hasil prediksi LightGBM"):
            lgb_show = test_out[['Tahun', 'Bulan', 'Jenis Kopi', 'Total Pendapatan (Rp)', 'Prediksi LightGBM']].copy()
            lgb_show['Total Pendapatan (Rp)'] = lgb_show['Total Pendapatan (Rp)'].apply(rupiah)
            lgb_show['Prediksi LightGBM'] = lgb_show['Prediksi LightGBM'].apply(rupiah)
            st.dataframe(lgb_show, use_container_width=True, hide_index=True)

    # ---------- Sub-tab: Perbandingan Performa ----------
    with subtab_perf:
        st.markdown("#### Perbandingan Performa Model")
        res_df = pd.DataFrame(results).T
        res_df.columns = ['MAE', 'RMSE', 'R²', 'Training Time (detik)']
        best_model = res_df['MAE'].idxmin()
        st.dataframe(res_df.style.format({'MAE': '{:,.0f}', 'RMSE': '{:,.0f}', 'R²': '{:.4f}', 'Training Time (detik)': '{:.4f}'}),
                     use_container_width=True)
        st.success(f"Model dengan performa terbaik (MAE terendah): **{best_model}**")

        st.markdown("#### Aktual vs Prediksi — Kedua Model (Data Uji)")
        jenis_perf = st.selectbox("Pilih jenis kopi", sorted(test_out['Jenis Kopi'].unique()), key="jenis_perf")
        sub_perf = test_out[test_out['Jenis Kopi'] == jenis_perf].copy()
        sub_perf['label'] = sub_perf['Bulan'].str[:3] + " " + sub_perf['Tahun'].astype(str)

        fig_perf = go.Figure()
        fig_perf.add_trace(go.Scatter(x=sub_perf['label'], y=sub_perf['Total Pendapatan (Rp)'], name="Aktual",
                                       mode='lines+markers', line=dict(color="#333", width=3)))
        fig_perf.add_trace(go.Scatter(x=sub_perf['label'], y=sub_perf['Prediksi Random Forest'], name="Prediksi RF",
                                       mode='lines+markers', line=dict(color=PRIMARY, dash='dash')))
        fig_perf.add_trace(go.Scatter(x=sub_perf['label'], y=sub_perf['Prediksi LightGBM'], name="Prediksi LightGBM",
                                       mode='lines+markers', line=dict(color=ACCENT, dash='dot')))
        fig_perf.update_layout(height=420, plot_bgcolor="white", yaxis_title="Pendapatan (Rp)")
        st.plotly_chart(fig_perf, use_container_width=True)

        with st.expander("Lihat tabel hasil prediksi lengkap (kedua model)"):
            test_show = test_out.copy()
            for col in ['Total Pendapatan (Rp)', 'Prediksi Random Forest', 'Prediksi LightGBM']:
                test_show[col] = test_show[col].apply(rupiah)
            st.dataframe(test_show, use_container_width=True, hide_index=True)

with tab4:
    st.subheader("Feature Importance — Random Forest vs LightGBM")
    fi_sorted = fi.sort_values('Random Forest', ascending=True)
    fig4 = go.Figure()
    fig4.add_trace(go.Bar(y=fi_sorted['Fitur'], x=fi_sorted['Random Forest'], name='Random Forest', orientation='h', marker_color=PRIMARY))
    fig4.add_trace(go.Bar(y=fi_sorted['Fitur'], x=fi_sorted['LightGBM'], name='LightGBM', orientation='h', marker_color=ACCENT))
    fig4.update_layout(height=420, barmode='group', plot_bgcolor="white", xaxis_title="Tingkat Kepentingan (dinormalisasi)")
    st.plotly_chart(fig4, use_container_width=True)
    st.caption("Fitur `lag_1` (pendapatan bulan sebelumnya) dan `jenis_kopi_enc` konsisten menjadi variabel paling berpengaruh pada kedua model.")

with tab5:
    st.subheader(f"Prediksi Pendapatan {next_bulan_nama} {next_tahun}")
    st.caption("Model dilatih ulang menggunakan seluruh data historis (Januari 2023 – bulan terakhir data) agar prediksi memanfaatkan informasi terbaru yang tersedia.")

    total_rf = forecast_df['Prediksi Random Forest (Rp)'].sum()
    total_lgb = forecast_df['Prediksi LightGBM (Rp)'].sum()

    c1, c2, c3 = st.columns(3)
    c1.metric(f"Total Prediksi {next_bulan_nama} {next_tahun} (Random Forest)", rupiah(total_rf))
    c2.metric(f"Total Prediksi {next_bulan_nama} {next_tahun} (LightGBM)", rupiah(total_lgb))
    c3.metric("Rata-rata Pendapatan Historis / Bulan", rupiah(avg_overall))

    st.subheader("Rincian Prediksi per Jenis Kopi")
    show_df = forecast_df.copy()
    for col in ['Pendapatan Bulan Terakhir (Rp)', 'Prediksi Random Forest (Rp)', 'Prediksi LightGBM (Rp)']:
        show_df[col] = show_df[col].apply(rupiah)
    st.dataframe(show_df, use_container_width=True, hide_index=True)

    fig5 = go.Figure()
    fig5.add_trace(go.Bar(x=forecast_df['Jenis Kopi'], y=forecast_df['Prediksi Random Forest (Rp)'], name='Prediksi RF', marker_color=PRIMARY))
    fig5.add_trace(go.Bar(x=forecast_df['Jenis Kopi'], y=forecast_df['Prediksi LightGBM (Rp)'], name='Prediksi LightGBM', marker_color=ACCENT))
    fig5.update_layout(height=420, barmode='group', plot_bgcolor="white", yaxis_title="Prediksi Pendapatan (Rp)")
    st.plotly_chart(fig5, use_container_width=True)

    st.info("Catatan asumsi: harga rata-rata memakai rata-rata 3 bulan terakhir per jenis kopi, dan kategori tren memakai kategori bulan terakhir yang datanya tersedia (karena kategori tren bulan depan belum bisa diketahui sebelum pendapatan aktualnya terjadi).")

with tab6:
    st.subheader("Laporan Penjualan")
    st.caption(
        "Pilih filter jenis kopi dan rentang periode, lalu unduh ringkasan laporan dalam format Word atau PDF. "
        "(Data sumber berupa transaksi per bulan, sehingga filter periode di sini berbentuk rentang bulan, bukan tanggal harian.)"
    )

    daily_periode = daily.copy()
    daily_periode['bulan_num'] = daily_periode['Bulan'].map(BULAN_MAP)
    daily_periode['periode'] = daily_periode['Tahun'] * 100 + daily_periode['bulan_num']

    jenis_list = sorted(daily_periode['Jenis Kopi'].unique())
    periode_sorted = sorted(daily_periode['periode'].unique())
    periode_label = {p: f"{BULAN_ORDER[(p % 100) - 1]} {p // 100}" for p in periode_sorted}

    jenis_filter = st.multiselect("Pilih jenis kopi", options=jenis_list, default=jenis_list)

    colp1, colp2 = st.columns(2)
    with colp1:
        periode_awal = st.selectbox("Dari bulan", periode_sorted, format_func=lambda p: periode_label[p],
                                     index=0, key="lap_periode_awal")
    with colp2:
        periode_akhir = st.selectbox("Sampai bulan", periode_sorted, format_func=lambda p: periode_label[p],
                                      index=len(periode_sorted) - 1, key="lap_periode_akhir")

    if periode_awal > periode_akhir:
        st.error("Bulan awal tidak boleh setelah bulan akhir. Silakan sesuaikan pilihan di atas.")
    elif not jenis_filter:
        st.warning("Pilih minimal satu jenis kopi untuk menampilkan laporan.")
    else:
        mask = (
            (daily_periode['periode'] >= periode_awal)
            & (daily_periode['periode'] <= periode_akhir)
            & (daily_periode['Jenis Kopi'].isin(jenis_filter))
        )
        df_filtered = daily_periode[mask].copy()

        if df_filtered.empty:
            st.warning("Tidak ada data untuk kombinasi filter yang dipilih.")
        else:
            total_pendapatan = df_filtered['Pendapatan (Rp)'].sum()
            total_qty = df_filtered['Jumlah Terjual'].sum()
            rata_harga = (
                (df_filtered['Harga (Rp)'] * df_filtered['Jumlah Terjual']).sum() / total_qty
                if total_qty > 0 else 0
            )

            st.markdown("#### Ringkasan Penjualan")
            c1, c2, c3 = st.columns(3)
            c1.metric("Total Pendapatan", rupiah(total_pendapatan))
            c2.metric("Total Unit Terjual", f"{int(total_qty):,}".replace(",", "."))
            c3.metric("Rata-rata Harga", rupiah(rata_harga))

            ringkasan_jenis = (
                df_filtered.groupby('Jenis Kopi')
                .agg(**{'Total Pendapatan (Rp)': ('Pendapatan (Rp)', 'sum'),
                        'Jumlah Terjual': ('Jumlah Terjual', 'sum')})
                .reset_index()
                .sort_values('Total Pendapatan (Rp)', ascending=False)
            )

            ringkasan_show = ringkasan_jenis.copy()
            ringkasan_show['Total Pendapatan (Rp)'] = ringkasan_show['Total Pendapatan (Rp)'].apply(rupiah)
            st.dataframe(ringkasan_show, use_container_width=True, hide_index=True)

            fig_lap = px.bar(ringkasan_jenis, x='Jenis Kopi', y='Total Pendapatan (Rp)',
                              color_discrete_sequence=[PRIMARY])
            fig_lap.update_layout(height=360, plot_bgcolor="white")
            st.plotly_chart(fig_lap, use_container_width=True)

            st.markdown("#### Unduh Laporan")
            periode_text = f"{periode_label[periode_awal]} – {periode_label[periode_akhir]}"
            jenis_text = "Semua Jenis Kopi" if len(jenis_filter) == len(jenis_list) else ", ".join(jenis_filter)
            file_tag = f"{periode_awal}_{periode_akhir}"

            colw, colpdf = st.columns(2)
            with colw:
                if DOCX_OK:
                    docx_bytes = generate_laporan_docx(periode_text, jenis_text, total_pendapatan,
                                                         total_qty, rata_harga, ringkasan_jenis)
                    st.download_button(
                        "⬇️ Unduh Laporan (Word)", data=docx_bytes,
                        file_name=f"laporan_penjualan_{file_tag}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    st.warning("Modul `python-docx` belum terpasang. Tambahkan `python-docx` ke requirements.txt.")
            with colpdf:
                if REPORTLAB_OK:
                    pdf_bytes = generate_laporan_pdf(periode_text, jenis_text, total_pendapatan,
                                                       total_qty, rata_harga, ringkasan_jenis)
                    st.download_button(
                        "⬇️ Unduh Laporan (PDF)", data=pdf_bytes,
                        file_name=f"laporan_penjualan_{file_tag}.pdf",
                        mime="application/pdf",
                    )
                else:
                    st.warning("Modul `reportlab` belum terpasang. Tambahkan `reportlab` ke requirements.txt.")

st.divider()
st.caption("Dashboard ini dijalankan di Google Colab menggunakan Python, Streamlit, Scikit-learn, dan LightGBM ")