"""
Modul logika data untuk dashboard Sumatra Roastery.
Berisi: pemuatan & pembentukan dataset, training model (Random Forest & LightGBM),
pembuatan laporan (Excel/Word/PDF), dan prediksi bulan berikutnya.
Dipisah dari app.py supaya app.py fokus ke tampilan (UI) saja.
"""
import streamlit as st
import pandas as pd
import numpy as np
import time
import io
import copy
import calendar as cal_module
from sklearn.ensemble import RandomForestRegressor
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
import lightgbm as lgb

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, HRFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False


# --- Identitas usaha untuk kop surat laporan (sesuaikan di sini jika berubah) ---
BUSINESS_NAME = "SUMATRA ROASTERY"
BUSINESS_TAGLINE = "Toko Bubuk Kopi"
BUSINESS_ADDRESS = "Jl. Ring Road No. 109, Sei Sikambing B, Kec. Medan Sunggal, Kota Medan, Sumatera Utara 20122"
BUSINESS_PHONE = "0811-610-088"
BUSINESS_WEBSITE = "sumatraroastery.com"
SIGNER_NAME = "Rahmad Fikri"
SIGNER_TITLE = "Pengelola Sumatra Roastery Medan"


def _docx_heading(doc, text, level=1):
    """Tambahkan heading dengan warna hijau brand (bukan biru bawaan Word),
    supaya senada dengan warna kop surat."""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(0x0F, 0x6B, 0x5C)
    return h


def _docx_add_bottom_border(paragraph, color="0F6B5C", size=12):
    """Tambahkan garis horizontal (border bawah) pada sebuah paragraph python-docx."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _docx_force_left(paragraph):
    """Pastikan paragraf benar-benar rata kiri tanpa indent, apa pun style bawaannya."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(0)


def _docx_float_logo_left(run, wrap_height_inches=0.95):
    """Ubah gambar inline (hasil run.add_picture) menjadi gambar melayang (anchored)
    yang menempel di kiri, dengan teks di paragraf-paragraf berikutnya melingkar
    di sebelah kanannya. Ini memungkinkan teks nama usaha & alamat tampil rata
    tengah di ruang sebelah kanan logo, seperti kop surat resmi."""
    drawing = run._r.find(qn('w:drawing'))
    if drawing is None:
        return
    inline = drawing.find(qn('wp:inline'))
    if inline is None:
        return

    extent = inline.find(qn('wp:extent'))
    doc_pr = inline.find(qn('wp:docPr'))
    cnv_pr = inline.find(qn('wp:cNvGraphicFramePr'))
    graphic = inline.find(qn('a:graphic'))

    cy = extent.get('cy') if extent is not None else str(int(Inches(0.6)))

    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '137160')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '251659264')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    anchor.append(simple_pos)

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'column')
    off_h = OxmlElement('wp:posOffset')
    off_h.text = '0'
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'paragraph')
    off_v = OxmlElement('wp:posOffset')
    off_v.text = '0'
    pos_v.append(off_v)
    anchor.append(pos_v)

    if extent is not None:
        anchor.append(copy.deepcopy(extent))
    eff_extent = OxmlElement('wp:effectExtent')
    eff_extent.set('l', '0'); eff_extent.set('t', '0'); eff_extent.set('r', '0'); eff_extent.set('b', '0')
    anchor.append(eff_extent)

    wrap_square = OxmlElement('wp:wrapSquare')
    wrap_square.set('wrapText', 'right')
    anchor.append(wrap_square)

    if doc_pr is not None:
        anchor.append(copy.deepcopy(doc_pr))
    if cnv_pr is not None:
        anchor.append(copy.deepcopy(cnv_pr))
    if graphic is not None:
        anchor.append(copy.deepcopy(graphic))

    drawing.remove(inline)
    drawing.append(anchor)


def _docx_add_letterhead(doc, logo_path=None):
    """Tambahkan kop surat (logo + identitas usaha) di bagian atas dokumen Word.
    Logo melayang (anchored) di kiri, nama usaha & alamat rata tengah di ruang
    sebelah kanannya — meniru format kop surat resmi instansi."""
    logo_inserted = False
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)

    if logo_path:
        try:
            run_logo = p_name.add_run()
            run_logo.add_picture(logo_path, width=Inches(0.85))
            _docx_float_logo_left(run_logo)
            logo_inserted = True
        except Exception:
            pass
    if not logo_inserted:
        try:
            st.warning(
                "Logo tidak ditemukan/gagal dimuat untuk laporan — pastikan file `logo.png` "
                "berada satu folder dengan app.py. Laporan tetap dibuat tanpa logo."
            )
        except Exception:
            pass

    run_name = p_name.add_run(f"{BUSINESS_NAME} - {BUSINESS_TAGLINE}")
    run_name.bold = True
    run_name.font.size = Pt(15)
    run_name.font.color.rgb = RGBColor(0x0F, 0x6B, 0x5C)

    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.paragraph_format.space_after = Pt(10)
    contact_line = f"{BUSINESS_ADDRESS}   |   Telepon: {BUSINESS_PHONE}   |   {BUSINESS_WEBSITE}"
    run_addr = p_addr.add_run(contact_line)
    run_addr.font.size = Pt(9.5)
    run_addr.font.color.rgb = RGBColor(0x3B, 0x2A, 0x20)

    line_p = doc.add_paragraph()
    _docx_force_left(line_p)
    _docx_add_bottom_border(line_p)
    doc.add_paragraph()


def _docx_add_signature(doc):
    """Tambahkan blok tanda tangan pemilik di bagian bawah dokumen Word."""
    SIGN_INDENT = Inches(3.0)

    doc.add_paragraph()
    p_place = doc.add_paragraph()
    p_place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_place.paragraph_format.left_indent = SIGN_INDENT
    p_place.add_run(f"Medan, {pd.Timestamp.now().strftime('%d %B %Y')}")

    for _ in range(3):
        doc.add_paragraph()

    p_signer_name = doc.add_paragraph()
    p_signer_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_signer_name.paragraph_format.left_indent = SIGN_INDENT
    run_signer_name = p_signer_name.add_run(SIGNER_NAME)
    run_signer_name.bold = True
    run_signer_name.underline = True

    p_signer_title = doc.add_paragraph()
    p_signer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_signer_title.paragraph_format.left_indent = SIGN_INDENT
    p_signer_title.add_run(SIGNER_TITLE)


BULAN_ORDER = ['Januari','Februari','Maret','April','Mei','Juni','Juli','Agustus','September','Oktober','November','Desember']
BULAN_MAP = {b: i + 1 for i, b in enumerate(BULAN_ORDER)}


@st.cache_data
def load_raw(file):
    xls = pd.ExcelFile(file)
    daily = pd.read_excel(xls, sheet_name='Dataset Harian 2023-2025', header=3)
    per_jenis = pd.read_excel(xls, sheet_name='Pendapatan Per Jenis Kopi', header=2)
    rekap = pd.read_excel(xls, sheet_name='Rekap Bulanan', header=2)
    return daily, per_jenis, rekap


@st.cache_data
def build_dataset(daily, per_jenis, rekap):
    agg_harga = (
        daily.groupby(['Tahun', 'Bulan', 'Jenis Kopi'])
        .apply(lambda g: (g['Harga (Rp)'] * g['Jumlah Terjual']).sum() / g['Jumlah Terjual'].sum()
               if g['Jumlah Terjual'].sum() > 0 else g['Harga (Rp)'].mean())
        .reset_index(name='harga_rata2')
    )

    df = per_jenis.merge(agg_harga, on=['Tahun', 'Bulan', 'Jenis Kopi'], how='left')
    df['bulan_num'] = df['Bulan'].map(BULAN_MAP)
    df['periode'] = df['Tahun'] * 100 + df['bulan_num']
    df = df.sort_values(['Jenis Kopi', 'periode']).reset_index(drop=True)
    df['lag_1'] = df.groupby('Jenis Kopi')['Total Pendapatan (Rp)'].shift(1)

    rekap = rekap.copy()
    rekap['bulan_num'] = rekap['Bulan'].map(BULAN_MAP)
    rekap['periode'] = rekap['Tahun'] * 100 + rekap['bulan_num']
    avg_overall = rekap['Total Pendapatan (Rp)'].mean()
    rekap['kategori_tren'] = np.where(rekap['Total Pendapatan (Rp)'] >= avg_overall, 'Tinggi', 'Rendah')

    df = df.merge(rekap[['periode', 'kategori_tren']], on='periode', how='left')
    return df, rekap, avg_overall


@st.cache_data
def train_models(df, split_ratio):
    le_jenis = LabelEncoder()
    le_tren = LabelEncoder()
    df_model = df.dropna(subset=['lag_1']).copy()
    df_model['jenis_kopi_enc'] = le_jenis.fit_transform(df_model['Jenis Kopi'])
    df_model['kategori_tren_enc'] = le_tren.fit_transform(df_model['kategori_tren'])
    df_model = df_model.sort_values('periode').reset_index(drop=True)

    features = ['Tahun', 'bulan_num', 'jenis_kopi_enc', 'harga_rata2', 'lag_1', 'kategori_tren_enc']
    target = 'Total Pendapatan (Rp)'

    periodes = sorted(df_model['periode'].unique())
    split_periode = periodes[int(len(periodes) * split_ratio)]
    train = df_model[df_model['periode'] < split_periode]
    test = df_model[df_model['periode'] >= split_periode]

    X_train, y_train = train[features], train[target]
    X_test, y_test = test[features], test[target]

    results = {}
    preds = {}

    t0 = time.time()
    rf = RandomForestRegressor(n_estimators=200, random_state=42)
    rf.fit(X_train, y_train)
    rf_time = time.time() - t0
    pred_rf = rf.predict(X_test)
    results['Random Forest'] = dict(
        MAE=mean_absolute_error(y_test, pred_rf),
        RMSE=np.sqrt(mean_squared_error(y_test, pred_rf)),
        R2=r2_score(y_test, pred_rf),
        Training_Time=rf_time,
    )
    preds['Random Forest'] = pred_rf

    t0 = time.time()
    lgbm = lgb.LGBMRegressor(n_estimators=200, random_state=42, verbose=-1)
    lgbm.fit(X_train, y_train)
    lgb_time = time.time() - t0
    pred_lgb = lgbm.predict(X_test)
    results['LightGBM'] = dict(
        MAE=mean_absolute_error(y_test, pred_lgb),
        RMSE=np.sqrt(mean_squared_error(y_test, pred_lgb)),
        R2=r2_score(y_test, pred_lgb),
        Training_Time=lgb_time,
    )
    preds['LightGBM'] = pred_lgb

    fi = pd.DataFrame({
        'Fitur': features,
        'Random Forest': rf.feature_importances_ / rf.feature_importances_.sum(),
        'LightGBM': lgbm.feature_importances_ / lgbm.feature_importances_.sum(),
    })

    test_out = test[['Tahun', 'Bulan', 'Jenis Kopi', 'Total Pendapatan (Rp)']].copy()
    test_out['Prediksi Random Forest'] = pred_rf
    test_out['Prediksi LightGBM'] = pred_lgb

    return results, fi, test_out, split_periode


def rupiah(x):
    return f"Rp {x:,.0f}".replace(",", ".")


def to_excel_bytes(df, sheet_name="Data"):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return buf.getvalue()


def generate_laporan_docx(periode_text, jenis_text, total_pendapatan, total_qty, rata_harga,
                           ringkasan_jenis, logo_path=None):
    doc = Document()
    _docx_add_letterhead(doc, logo_path=logo_path)

    _docx_heading(doc, 'Laporan Penjualan', level=1)
    doc.add_paragraph(f"Periode: {periode_text}")
    doc.add_paragraph(f"Jenis Kopi: {jenis_text}")
    doc.add_paragraph(f"Tanggal Cetak: {pd.Timestamp.now().strftime('%d %B %Y')}")

    _docx_heading(doc, 'Ringkasan', level=2)
    doc.add_paragraph(f"Total Pendapatan: {rupiah(total_pendapatan)}")
    doc.add_paragraph(f"Total Unit Terjual: {int(total_qty):,}".replace(",", "."))
    doc.add_paragraph(f"Rata-rata Harga: {rupiah(rata_harga)}")

    _docx_heading(doc, 'Rincian per Jenis Kopi', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Jenis Kopi', 'Total Pendapatan (Rp)', 'Jumlah Terjual'
    for _, row in ringkasan_jenis.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Jenis Kopi'])
        cells[1].text = rupiah(row['Total Pendapatan (Rp)'])
        cells[2].text = str(int(row['Jumlah Terjual']))

    _docx_add_signature(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_letterhead_elements(styles, logo_path=None):
    """Buat elemen kop surat (logo + identitas usaha + garis) untuk laporan PDF.
    Teks identitas dibuat RATA TENGAH sebagai satu blok (bukan rata kiri) di ruang
    sebelah kanan logo, meniru format kop surat resmi instansi (judul, subjudul,
    alamat, telepon semuanya center sehingga tepinya 'naik turun')."""
    name_style = ParagraphStyle('BizName', parent=styles['Normal'], fontSize=15, leading=18,
                                 textColor=colors.HexColor("#0F6B5C"), fontName='Helvetica-Bold',
                                 alignment=TA_CENTER)
    info_style = ParagraphStyle('BizInfo', parent=styles['Normal'], fontSize=9.5, leading=13,
                                 textColor=colors.HexColor("#3B2A20"), alignment=TA_CENTER)

    contact_line = f"{BUSINESS_ADDRESS}   |   Telepon: {BUSINESS_PHONE}   |   {BUSINESS_WEBSITE}"
    text_cell = [
        Paragraph(f"{BUSINESS_NAME} - {BUSINESS_TAGLINE}", name_style),
        Paragraph(contact_line, info_style),
    ]

    if logo_path:
        try:
            logo = RLImage(logo_path, width=22 * mm, height=22 * mm)
        except Exception:
            logo = ""
    else:
        logo = ""

    # Kolom teks dibuat lebih lebar & kolom logo tetap kecil di kiri, supaya blok
    # teks yang center benar-benar center relatif ke sisa lebar halaman (bukan
    # ketarik ke kiri oleh kolom logo yang sempit).
    head_table = Table([[logo, text_cell]], colWidths=[24 * mm, 142 * mm])
    head_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, 0), 'CENTER'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (0, 0), (0, 0), 4),
        ('LEFTPADDING', (1, 0), (1, 0), 0),
        ('RIGHTPADDING', (1, 0), (1, 0), 0),
    ]))

    return [
        head_table,
        Spacer(1, 6),
        HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#0F6B5C")),
        Spacer(1, 14),
    ]


def _pdf_signature_elements(styles):
    """Buat elemen blok tanda tangan pemilik untuk laporan PDF."""
    SIGN_INDENT = 240
    center_style = ParagraphStyle('SignCenter', parent=styles['Normal'], alignment=TA_CENTER,
                                   fontSize=10, leftIndent=SIGN_INDENT)
    center_bold = ParagraphStyle('SignCenterBold', parent=center_style, fontName='Helvetica-Bold')
    return [
        Spacer(1, 24),
        Paragraph(f"Medan, {pd.Timestamp.now().strftime('%d %B %Y')}", center_style),
        Spacer(1, 42),
        Paragraph(f"<u>{SIGNER_NAME}</u>", center_bold),
        Paragraph(SIGNER_TITLE, center_style),
    ]


def generate_laporan_pdf(periode_text, jenis_text, total_pendapatan, total_qty, rata_harga,
                          ringkasan_jenis, logo_path=None):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = _pdf_letterhead_elements(styles, logo_path=logo_path)
    elements += [
        Paragraph("Laporan Penjualan", styles['Title']),
        Spacer(1, 12),
        Paragraph(f"Periode: {periode_text}", styles['Normal']),
        Paragraph(f"Jenis Kopi: {jenis_text}", styles['Normal']),
        Paragraph(f"Tanggal Cetak: {pd.Timestamp.now().strftime('%d %B %Y')}", styles['Normal']),
        Spacer(1, 12),
        Paragraph("Ringkasan", styles['Heading2']),
        Paragraph(f"Total Pendapatan: {rupiah(total_pendapatan)}", styles['Normal']),
        Paragraph(f"Total Unit Terjual: {int(total_qty):,}".replace(",", "."), styles['Normal']),
        Paragraph(f"Rata-rata Harga: {rupiah(rata_harga)}", styles['Normal']),
        Spacer(1, 12),
        Paragraph("Rincian per Jenis Kopi", styles['Heading2']),
        Spacer(1, 6),
    ]

    data = [['Jenis Kopi', 'Total Pendapatan (Rp)', 'Jumlah Terjual']]
    for _, row in ringkasan_jenis.iterrows():
        data.append([row['Jenis Kopi'], rupiah(row['Total Pendapatan (Rp)']), str(int(row['Jumlah Terjual']))])
    tbl = Table(data, hAlign='LEFT')
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0F6B5C")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
    ]))
    elements.append(tbl)
    elements += _pdf_signature_elements(styles)
    doc.build(elements)
    return buf.getvalue()


@st.cache_data
def forecast_next_month(df):
    """Latih model pakai SELURUH data (bukan hanya data training) lalu prediksi bulan setelah periode terakhir."""
    le_jenis = LabelEncoder()
    le_tren = LabelEncoder()
    df_model = df.dropna(subset=['lag_1']).copy()
    df_model['jenis_kopi_enc'] = le_jenis.fit_transform(df_model['Jenis Kopi'])
    df_model['kategori_tren_enc'] = le_tren.fit_transform(df_model['kategori_tren'])
    df_model = df_model.sort_values('periode').reset_index(drop=True)

    features = ['Tahun', 'bulan_num', 'jenis_kopi_enc', 'harga_rata2', 'lag_1', 'kategori_tren_enc']
    target = 'Total Pendapatan (Rp)'

    rf_full = RandomForestRegressor(n_estimators=200, random_state=42)
    rf_full.fit(df_model[features], df_model[target])
    lgb_full = lgb.LGBMRegressor(n_estimators=200, random_state=42, verbose=-1)
    lgb_full.fit(df_model[features], df_model[target])

    last_periode = df_model['periode'].max()
    last_tahun, last_bulan_num = last_periode // 100, last_periode % 100
    next_bulan_num = 1 if last_bulan_num == 12 else last_bulan_num + 1
    next_tahun = last_tahun + 1 if last_bulan_num == 12 else last_tahun
    next_bulan_nama = BULAN_ORDER[next_bulan_num - 1]

    last_kategori_tren = df_model.loc[df_model['periode'] == last_periode, 'kategori_tren_enc'].iloc[0]

    rows = []
    for jenis in df_model['Jenis Kopi'].unique():
        sub = df_model[df_model['Jenis Kopi'] == jenis].sort_values('periode')
        lag_1_next = sub[target].iloc[-1]
        harga_next = sub['harga_rata2'].tail(3).mean()
        jenis_enc = sub['jenis_kopi_enc'].iloc[-1]
        x_next = pd.DataFrame([{
            'Tahun': next_tahun, 'bulan_num': next_bulan_num, 'jenis_kopi_enc': jenis_enc,
            'harga_rata2': harga_next, 'lag_1': lag_1_next, 'kategori_tren_enc': last_kategori_tren,
        }])[features]
        pred_rf = rf_full.predict(x_next)[0]
        pred_lgb = lgb_full.predict(x_next)[0]
        rows.append({'Jenis Kopi': jenis, 'Pendapatan Bulan Terakhir (Rp)': lag_1_next,
                      'Prediksi Random Forest (Rp)': pred_rf, 'Prediksi LightGBM (Rp)': pred_lgb})

    forecast_df = pd.DataFrame(rows).sort_values('Prediksi Random Forest (Rp)', ascending=False).reset_index(drop=True)
    return forecast_df, next_bulan_nama, next_tahun


# ---------------------------------------------------------------------------
# KALENDER & ESTIMASI HARIAN/MINGGUAN
# Catatan penting: dataset sumber hanya mencatat transaksi pada level
# Tahun + Bulan (tidak ada tanggal harian), begitu pula model Random Forest
# dan LightGBM dilatih dan dievaluasi pada level bulanan (lihat BAB III).
# Karena itu, nilai "harian" dan "mingguan" di bagian ini BUKAN hasil
# prediksi model, melainkan estimasi kasar dari pembagian rata pendapatan
# bulanan (aktual atau prediksi) sesuai jumlah hari/minggu kalender pada
# bulan tersebut. Tujuannya murni membantu visualisasi & perencanaan,
# bukan mengklaim akurasi prediksi pada level harian/mingguan.
# ---------------------------------------------------------------------------

def build_calendar_matrix(rekap):
    """Bentuk matriks Tahun x Bulan berisi Total Pendapatan (Rp) aktual,
    dipakai untuk kalender heatmap tahunan."""
    pivot = rekap.pivot_table(index='Tahun', columns='Bulan', values='Total Pendapatan (Rp)', aggfunc='sum')
    pivot = pivot.reindex(columns=BULAN_ORDER)
    pivot = pivot.sort_index()
    return pivot


def get_periode_options(rekap, forecast_total, next_bulan_nama, next_tahun):
    """Daftar (Tahun, Bulan) yang bisa dipilih di kalender: seluruh bulan
    aktual pada data + satu bulan prediksi (bulan berikutnya setelah data
    terakhir)."""
    opsi = [
        {'Tahun': int(row['Tahun']), 'Bulan': row['Bulan'], 'bulan_num': BULAN_MAP[row['Bulan']],
         'Total Pendapatan (Rp)': row['Total Pendapatan (Rp)'], 'tipe': 'Aktual'}
        for _, row in rekap.iterrows()
    ]
    opsi.append({
        'Tahun': int(next_tahun), 'Bulan': next_bulan_nama, 'bulan_num': BULAN_MAP[next_bulan_nama],
        'Total Pendapatan (Rp)': forecast_total, 'tipe': 'Prediksi',
    })
    return opsi


def daily_weekly_estimate(tahun, bulan_num, total_pendapatan):
    """Bagi rata total pendapatan satu bulan ke estimasi harian, dan ke
    estimasi mingguan mengikuti struktur kalender (Senin-Minggu) bulan itu."""
    _, num_days = cal_module.monthrange(int(tahun), int(bulan_num))
    daily_avg = total_pendapatan / num_days

    weeks = cal_module.Calendar(firstweekday=0).monthdatescalendar(int(tahun), int(bulan_num))
    rows = []
    week_no = 1
    for week in weeks:
        days_in_month = [d for d in week if d.month == int(bulan_num)]
        if not days_in_month:
            continue
        n = len(days_in_month)
        rows.append({
            'Minggu ke': week_no,
            'Rentang Tanggal': f"{days_in_month[0].strftime('%d %b')} \u2013 {days_in_month[-1].strftime('%d %b')}",
            'Jumlah Hari': n,
            'Estimasi Pendapatan (Rp)': daily_avg * n,
        })
        week_no += 1

    weekly_df = pd.DataFrame(rows)
    return num_days, daily_avg, weekly_df